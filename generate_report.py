import sys
import subprocess
import os

def install_and_import(package, import_name=None):
    if import_name is None:
        import_name = package
    try:
        __import__(import_name)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

install_and_import('python-docx', 'docx')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document()
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

doc.add_heading('微细管状网络分割提分与拓扑保全工程化方案报告', 0)

doc.add_heading('一、 核心观点与总体架构', level=1)
doc.add_paragraph('核心矛盾：管状结构分割的核心矛盾在于“极端类别不平衡导致的背景主导”与“体素级 Dice 无法惩罚拓扑断裂”。').style = 'List Bullet'
doc.add_paragraph('提分关键：推行“Patch 前景过采样 + 拓扑感知网络（条形池化/可变形卷积）+ 复合拓扑损失（clDice + Focal Tversky）+ 双阈值迟滞后处理”的工程组合拳。').style = 'List Bullet'

doc.add_heading('二、 逐步论证与核心方案设计', level=1)

doc.add_heading('1. 复合损失函数体系（解决断裂与极度不平衡）', level=2)
doc.add_paragraph('常规 Dice Loss 仅关注体素重合面积，对于仅占 1–2 个像素宽度的微细末梢分支断裂不敏感。建议构建三元联合损失函数：\nL_total = λ1 * L_Focal + λ2 * L_Tversky + λ3 * L_clDice')
doc.add_paragraph('Focal Tversky Loss（压制假阴性）：设置 α = 0.3, β = 0.7。提高 β 权重可严厉惩罚漏检（FN），迫使模型强行关注并召回低对比度末梢。').style = 'List Bullet'
doc.add_paragraph('clDice（Centerline Dice Loss - 拓扑保全核心）：利用可微分的软骨架化算法（Soft-Skeletonization）提取预测与标注的中心线，计算拓扑重合度。该损失直接惩罚血管断裂，可显著改善连通性。').style = 'List Bullet'
doc.add_paragraph('【补充优化】：clDice 在 3D 数据上的计算非常消耗显存，建议在训练最后 30% Epoch 介入微调。此外，可探索基于持续同调的 TopoLoss 或 Betti Matching 作为更严谨的拓扑损失平替。').style = 'List Bullet'

doc.add_heading('2. 网络架构选型与针对性改造（解决空间细节丢失）', level=2)
doc.add_paragraph('避免暴力下采样：优先采用 nnU-Net 框架（针对体素间距自适应调整 Pooling 次数）或 High-Resolution Network (HRNet) 维持高分辨率主干。').style = 'List Bullet'
doc.add_paragraph('长程管状特征捕获（Strip Pooling）：用条形池化替换标准方形池化，匹配血管沿特定几何方向延伸的物理特征。').style = 'List Bullet'
doc.add_paragraph('可变形卷积（DCNv2/v3）：在编码器深层引入 Deformable Convolution，让感受野根据血管弯曲走向自适应形变，精准贴合管壁曲率。').style = 'List Bullet'
doc.add_paragraph('【补充优化】：针对长程拓扑连续性，可考虑引入状态空间模型（SSMs），例如 Vision Mamba (如 U-Mamba)。其处理长序列时显存占用远小于 Transformer，且在医学管状结构分割上表现出极强的连通性保持能力。').style = 'List Bullet'

doc.add_heading('3. 数据预处理与动态采样策略（解决小目标丢失）', level=2)
doc.add_paragraph('前景过采样（Foreground Oversampling）：严禁整图暴力 Resize。采用 Sliding Patch 裁剪，训练时强制保证每个 Batch 中至少 60%–70% 的 Patch 包含血管目标，消除纯背景空 Patch 对梯度的稀释。').style = 'List Bullet'
doc.add_paragraph('对比度局部自适应增强（CLAHE）：在数据流水线中引入自适应直方图均衡化，拉开低信噪比末梢与背景的灰度差异。').style = 'List Bullet'
doc.add_paragraph('【补充优化】：增加弹性形变（Elastic Deformation）作为核心数据增强方式，以模拟血管的自然扭曲，极大地扩充形态学多样性。').style = 'List Bullet'

doc.add_heading('4. 推理与精细后处理（解决孤立噪点与微小断点）', level=2)
doc.add_paragraph('测试时增强与高斯滑动窗口（TTA + Gaussian Sliding Window）：推理时 Patch 边缘施加高斯权重融合，配合多尺度与翻转 TTA，平滑边界过渡伪影。').style = 'List Bullet'
p = doc.add_paragraph('双阈值迟滞分割（Hysteresis Thresholding）：\n- 高阈值（如 0.6）：确定可靠的主干血管（低 FP）。\n- 低阈值（如 0.25）：捕获微弱末梢。\n- 从高阈值连通域出发，仅保留与主干相连的低阈值区域，滤除背景孤立噪点。')
p.style = 'List Bullet'
doc.add_paragraph('形态学连通域滤波：滤除像素体积小于特定阈值的微小离散假阳性连通块。').style = 'List Bullet'

doc.add_heading('三、 实用操作项与优先级建议', level=1)
doc.add_paragraph('【新增】P-1（指标基建）：在执行方案前，在验证集中加入 clDice、clF1 或 Error of Betti Number。拓扑改善在全局 Dice 上可能仅体现为微小提升，需依赖拓扑维度指标指导模型选择，避免误杀好模型。').style = 'List Bullet'
doc.add_paragraph('P0（立即见效 - 损失函数替换）：将现有损失替换为 0.5 * Focal_Loss + 0.5 * Tversky_Loss(alpha=0.3, beta=0.7)；训练中后期接入 clDice。').style = 'List Bullet'
doc.add_paragraph('P1（训练策略调整 - 采样与尺度）：切换为 Patch 级训练，Foreground Crop 比例 ≥ 0.6；加入 CLAHE 和弹性形变增强。').style = 'List Bullet'
doc.add_paragraph('P2（架构替换/优化）：以标准 nnU-Net 为基座，考虑加入 Strip Pooling、DCN 或探索 U-Mamba 架构。').style = 'List Bullet'
doc.add_paragraph('P3（推理后处理调优）：在验证集上扫描最优的 Hysteresis 双阈值组合，并执行连通域去噪。').style = 'List Bullet'

doc.add_heading('四、 限制与应对策略', level=1)
doc.add_paragraph('限制 1（算力与显存消耗）：clDice 提取 Soft-Skeleton 开销大。\n【应对方案】：显存紧张时，仅在训练最后 30% Epoch 介入微调。').style = 'List Bullet'
doc.add_paragraph('限制 2（标注噪声敏感性）：提高漏报惩罚时，若数据本身存在漏标，模型易将背景噪声误学为正例。\n【应对方案】：辅以微小的权重衰减（Weight Decay）。如果漏标严重，建议引入 Noise-robust Loss（如 Symmetric Cross Entropy），或在训练中后期引入伪标签（Pseudo-labeling）机制（将低阈值连通末梢作为正样本加入自训练）。').style = 'List Bullet'

doc.add_heading('五、 总结', level=1)
doc.add_paragraph('该路线结合 Focal Tversky 与 clDice 强化召回与连通性，辅以 Patch 前景过采样与双阈值迟滞后处理，在理论与工程上均具备极高的落地价值。配合良好的工程调参，将 Baseline 从 0.65 提升至 0.80+ 具备充足的可行性。')

output_path = r'c:\Users\12751\Downloads\DSCA数据集-已整合\微细管状网络分割优化方案报告.docx'
doc.save(output_path)
print(f"Document successfully saved to {output_path}")
